#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
处理docx文件，优化格式用于公众号导入
- 清理多余的换行
- 设置美观的字体
- 优化段落格式
- 智能合并短段落
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def is_title(text):
    """判断是否为标题"""
    text = text.strip()
    
    # 检查是否以emoji图标开头（常见标题特征）
    # emoji通常在文本开头，后面跟空格和标题文字
    emoji_pattern = re.compile(r'^[\U0001F300-\U0001F9FF\U0001FA00-\U0001FAFF\U00002600-\U000027BF\U0001F600-\U0001F64F\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF]+')
    if emoji_pattern.match(text):
        # 如果以emoji开头且长度较短，很可能是标题
        if len(text) < 50:
            return True
    
    # 移除emoji后检查标题
    text_without_emoji = re.sub(r'[\U0001F300-\U0001F9FF\U0001FA00-\U0001FAFF\U00002600-\U000027BF\U0001F600-\U0001F64F\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF]+', '', text).strip()
    
    # 常见标题列表（不含emoji）
    common_titles = [
        "引言", "项目概述", "设计哲学", "平台定位", "核心价值", "技术架构",
        "模块化设计理念", "数据流转架构", "存储方案", "核心AI能力", "全面的AI技术栈",
        "革命性的零样本标注技术", "多场景预训练模型", "IoT能力", "完整的设备生命周期管理",
        "强大的规则引擎", "数据智能分析", "部署灵活性", "独立部署优势", "一键部署方案",
        "应用场景适配", "核心优势", "多语言混编架构", "零样本标注技术", "云边端灵活部署",
        "丰富生态支持", "持续迭代优化", "应用场景", "人群密度管控", "周界防护",
        "跌倒检测", "异常逗留识别", "肢体冲突预警", "非法闯入检测", "公共场所控烟",
        "人流统计管控", "区域越界预警", "环境安全检查", "火灾早预警", "扩展应用领域",
        "系统展示", "核心功能界面展示", "技术实现", "设备控制核心逻辑", "安全认证体系",
        "AI模型管理", "高性能任务处理", "功能介绍", "设备管理模块", "流媒体管理模块",
        "数据标注模块", "模型训练与管理模块", "AI智能分析模块", "规则引擎模块",
        "系统管理模块", "数据统计与分析模块", "部署安装", "部署要求", "部署优势",
        "社区与开源", "我们的承诺", "加入我们", "演示环境与支持", "在线演示",
        "结语", "联系方式"
    ]
    
    # 检查是否是常见标题（含或不含emoji）
    if text_without_emoji in common_titles or text in common_titles:
        return True
    
    # 检查是否包含常见标题关键词
    title_keywords = ['概述', '引言', '结语', '模块', '能力', '架构', '方案', '场景', 
                     '优势', '技术', '管理', '系统', '部署', '介绍', '展示', '实现']
    if any(keyword in text_without_emoji for keyword in title_keywords) and len(text) < 50:
        return True
    
    # 标题特征：短文本、可能包含数字编号、可能全角标点结尾
    if len(text) < 50 and (text.endswith('：') or text.endswith(':') or 
                          re.match(r'^[一二三四五六七八九十\d]+[、.．]', text_without_emoji) or
                          re.match(r'^第[一二三四五六七八九十\d]+[章节部分]', text_without_emoji)):
        return True
    
    return False

def is_project_url(text):
    """判断是否包含项目地址"""
    # 匹配项目地址模式
    url_pattern = r'项目地址[：:]\s*https?://[^\s]+'
    return bool(re.search(url_pattern, text))

def should_merge(prev_text, current_text):
    """判断是否应该合并段落"""
    # 如果前一个段落很短（少于30字）且不是标题，且当前段落也不是标题，可以合并
    if len(prev_text) < 30 and not is_title(prev_text) and not is_title(current_text):
        # 如果前一个段落不以句号、问号、感叹号结尾，可以合并
        if not prev_text.endswith(('。', '！', '？', '.', '!', '?')):
            return True
    return False

def process_docx(input_path, output_path=None):
    """
    处理docx文件，优化格式
    
    Args:
        input_path: 输入docx文件路径
        output_path: 输出docx文件路径，如果为None则自动生成
    """
    if not os.path.exists(input_path):
        print(f"错误：文件不存在: {input_path}")
        return False
    
    # 如果没有指定输出路径，自动生成
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}_优化版.docx"
    
    # 读取文档
    doc = Document(input_path)
    
    # 公众号常用字体：微软雅黑、思源黑体、PingFang SC等
    font_name = "微软雅黑"  # 公众号常用字体
    
    # 收集所有段落文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 清理段落内的多余空格
            text = re.sub(r' +', ' ', text)
            # 清理段落内的多余换行（保留单个换行用于列表等）
            text = re.sub(r'\n{2,}', '\n', text)
            paragraphs.append(text)
    
    # 智能合并段落（确保标题独立成行）
    merged_paragraphs = []
    for i, text in enumerate(paragraphs):
        if not merged_paragraphs:
            merged_paragraphs.append(text)
        else:
            prev_text = merged_paragraphs[-1]
            # 如果当前是标题，必须独立成行，不合并
            if is_title(text):
                merged_paragraphs.append(text)
            # 如果前一个是标题，当前不是标题，不合并（标题后应该是正文）
            elif is_title(prev_text):
                merged_paragraphs.append(text)
            # 如果应该合并，则合并
            elif should_merge(prev_text, text):
                merged_paragraphs[-1] = prev_text + text
            else:
                merged_paragraphs.append(text)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置默认字体
    style = new_doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(14)  # 公众号常用字号：14pt（更美观）
    
    # 添加处理后的段落
    for text in merged_paragraphs:
        # 如果段落中包含换行，可能是列表或需要保留的结构
        if '\n' in text:
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    para = new_doc.add_paragraph(line)
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)  # 段落间距（更紧凑）
                    para_format.line_spacing = 1.5  # 行距1.5倍（更美观）
                    
                    # 判断是否为标题，设置相应格式
                    if is_title(line):
                        para_format.space_before = Pt(12)  # 标题前间距
                        para_format.space_after = Pt(10)  # 标题后间距
                        # 清除原有runs，重新添加格式化的文本
                        para.clear()
                        run = para.add_run(line)
                        run.font.name = font_name
                        run.font.size = Pt(18)  # 标题18pt（更大更醒目）
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色（#0066CC）
                    elif is_project_url(line):
                        # 项目地址特殊处理：红色、小字体
                        para_format.space_after = Pt(6)
                        para.clear()
                        # 分离"项目地址："和URL
                        match = re.search(r'(项目地址[：:]\s*)(https?://[^\s]+)', line)
                        if match:
                            prefix = match.group(1)
                            url = match.group(2)
                            # 添加前缀（正常大小）
                            run1 = para.add_run(prefix)
                            run1.font.name = font_name
                            run1.font.size = Pt(12)  # 稍小字体
                            # 添加URL（红色、更小字体）
                            run2 = para.add_run(url)
                            run2.font.name = font_name
                            run2.font.size = Pt(11)  # 更小字体，不影响正文观感
                            run2.font.color.rgb = RGBColor(220, 20, 60)  # 红色（#DC143C）
                        else:
                            # 如果匹配失败，整个文本设为红色小字体
                            run = para.add_run(line)
                            run.font.name = font_name
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(220, 20, 60)
                    elif '项目地址' in line and 'http' in line:
                        # 行中包含项目地址但不是整行
                        para.clear()
                        url_pattern = r'(项目地址[：:]\s*)(https?://[^\s]+)'
                        parts = re.split(url_pattern, line)
                        for part in parts:
                            if not part:
                                continue
                            if re.match(r'https?://', part):
                                # URL部分：红色、小字体
                                run = para.add_run(part)
                                run.font.name = font_name
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(220, 20, 60)
                            elif part.startswith('项目地址'):
                                # "项目地址："前缀
                                run = para.add_run(part)
                                run.font.name = font_name
                                run.font.size = Pt(12)
                            else:
                                # 其他文本：正常格式
                                run = para.add_run(part)
                                run.font.name = font_name
                                run.font.size = Pt(14)
                    else:
                        # 普通正文
                        para.clear()
                        run = para.add_run(line)
                        run.font.name = font_name
                        run.font.size = Pt(14)  # 正文14pt（公众号常用）
        else:
            # 普通段落
            para = new_doc.add_paragraph(text)
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)  # 段落间距（更紧凑）
            para_format.line_spacing = 1.5  # 行距1.5倍（更美观）
            
            # 判断是否为标题
            if is_title(text):
                para_format.space_before = Pt(12)  # 标题前间距
                para_format.space_after = Pt(10)  # 标题后间距
                # 清除原有runs，重新添加格式化的文本
                para.clear()
                run = para.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(18)  # 标题18pt（更大更醒目）
                run.bold = True
                run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色（#0066CC）
            elif is_project_url(text):
                # 项目地址特殊处理：红色、小字体
                para_format.space_after = Pt(6)
                para.clear()
                # 分离"项目地址："和URL
                match = re.search(r'(项目地址[：:]\s*)(https?://[^\s]+)', text)
                if match:
                    prefix = match.group(1)
                    url = match.group(2)
                    # 添加前缀（正常大小）
                    run1 = para.add_run(prefix)
                    run1.font.name = font_name
                    run1.font.size = Pt(12)  # 稍小字体
                    # 添加URL（红色、更小字体）
                    run2 = para.add_run(url)
                    run2.font.name = font_name
                    run2.font.size = Pt(11)  # 更小字体，不影响正文观感
                    run2.font.color.rgb = RGBColor(220, 20, 60)  # 红色（#DC143C）
                else:
                    # 如果匹配失败，整个文本设为红色小字体
                    run = para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(220, 20, 60)
            else:
                # 普通正文，检查是否包含项目地址
                para.clear()
                # 检查段落中是否包含项目地址
                if '项目地址' in text and 'http' in text:
                    # 分段处理：项目地址部分用红色小字体，其他部分正常
                    url_pattern = r'(项目地址[：:]\s*)(https?://[^\s]+)'
                    parts = re.split(url_pattern, text)
                    for part in parts:
                        if not part:
                            continue
                        # 检查是否是URL部分
                        if re.match(r'https?://', part):
                            # URL部分：红色、小字体
                            run = para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(11)  # 更小字体
                            run.font.color.rgb = RGBColor(220, 20, 60)  # 红色
                        elif part.startswith('项目地址'):
                            # "项目地址："前缀：正常大小
                            run = para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(12)  # 稍小字体
                        else:
                            # 其他文本：正常格式
                            run = para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(14)  # 正文14pt
                else:
                    # 普通正文，直接设置格式
                    run = para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(14)  # 正文14pt（公众号常用）
    
    # 保存新文档
    new_doc.save(output_path)
    print(f"处理完成！已保存到: {output_path}")
    print(f"原段落数: {len(paragraphs)}, 优化后段落数: {len(merged_paragraphs)}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 process_docx.py <输入文件路径> [输出文件路径]")
        print("示例: python3 process_docx.py doc/项目宣传2_公众号版.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        process_docx(input_file, output_file)
    except ImportError:
        print("错误：需要安装python-docx库")
        print("请运行: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

